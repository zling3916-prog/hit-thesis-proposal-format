#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
hit-thesis-proposal-format: 哈尔滨工业大学硕士学位开题报告格式调整脚本

按照《硕士学位论文报告模板（2025）》的格式要求，把一份开题报告草稿 docx
统一为规范格式（页面设置 / 字体 / 层级标题 / 段落与行距 / 参考文献 / 封面）。

用法:
    python format_proposal.py <输入.docx> <输出.docx> [--cover-json 封面元数据.json] [--dry-run]

- 输入：任意内容完整的开题报告草稿（正文 + 层级标题 + 参考文献，可有封面）。
- 输出：套用模板格式后的新 docx；原文件不会被修改。
- --cover-json：可选。提供 JSON 时在文首重建模板封面，字段：
      {"title":"...","degree":"学位论文|实践成果","college":"...","discipline":"...",
       "advisor":"...","student":"...","student_id":"...","date":"..."}
- 仅改格式，不增删正文文字。

依赖: python-docx  (pip install python-docx)
"""
import argparse
import json
import os
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

# ---------- 常量：模板格式参数（与 references/format-spec.md 保持一致） ----------
PAGE = dict(width=Cm(21.0), height=Cm(29.7),
            top=Cm(2.5), bottom=Cm(2.2), left=Cm(2.5), right=Cm(2.5),
            header=Cm(1.5), footer=Cm(1.75))

LINE_TWIPS = 420          # 固定行距 21pt（w:line=420, lineRule=exact）

# 层级规则: 中文字体 / 字号pt / 段前行 / 段后行
LEVELS = {
    1: dict(east='黑体', size=15, before_lines=0.5, after_lines=0.5),   # 节
    2: dict(east='黑体', size=14, before_lines=0.5, after_lines=0.5),   # 条
    3: dict(east='黑体', size=12, before_lines=0,   after_lines=0),     # 款
    # 项（（1）…）：按用户要求与正文同格式，不再作为标题单独处理
}
BODY = dict(east='宋体', size=12, before_lines=0, after_lines=0)

LATIN_FONT = 'Times New Roman'
BODY_FIRST_TWIPS = 480    # 正文首行缩进2字符 ≈ 0.85cm
REF_HANG_TWIPS = 480      # 参考文献悬挂缩进2字符（左缩进+悬挂）
COVER_LEFT_TWIPS = 1260   # 封面填表行左缩进（2.25cm）
COVER_FILL_TAB_TWIPS = 4500  # 封面填表行填充区制表位（对齐下划线起点）

# 层级编号识别（同时兼容 1．/1. /1、 与 1.1、1.1.1、（1））
RE_L1 = re.compile(r'^\s*\d+\s*[．.、]\s*\S')
RE_L2 = re.compile(r'^\s*\d+\s*[．.]\s*\d+\s*\S')
RE_L3 = re.compile(r'^\s*\d+\s*[．.]\s*\d+\s*[．.]\s*\d+\s*\S')

# 封面行识别（限文档前 30 段）
COVER_PATTERNS = [
    ('school', re.compile(r'哈尔滨工业大学')),
    ('main_title', re.compile(r'硕士学位开题报告')),
    ('subtitle', re.compile(r'^[（(](学位论文|实践成果)[)）]')),
    ('topic', re.compile(r'^\s*题\s*目\s*[:：]')),
    ('college', re.compile(r'^\s*学\s*院\s*[（(]部[)）]')),
    ('discipline', re.compile(r'学科\s*/?\s*专业学位类别')),
    ('advisor', re.compile(r'^\s*导\s*师')),
    ('student', re.compile(r'^\s*研\s*究\s*生')),
    ('student_id', re.compile(r'^\s*学\s*号')),
    ('date', re.compile(r'开题报告日期')),
    ('foot', re.compile(r'研究生院制')),
]
COVER_MAX_INDEX = 30

# 封面填表行标签：用于把已有封面行切分为「标签(不加下划线) + 填充值(加下划线)」
FILL_LABELS = [
    (re.compile(r'^\s*学\s*院\s*[（(]部[)）]'), '学     院（部）'),
    (re.compile(r'学科\s*/?\s*专业学位类别'), '学科/专业学位类别'),
    (re.compile(r'^\s*导\s*师'), '导        师'),
    (re.compile(r'^\s*研\s*究\s*生'), '研   究   生'),
    (re.compile(r'^\s*学\s*号'), '学       号'),
    (re.compile(r'开题报告日期'), '开题报告日期'),
]

REF_HEAD_RE = re.compile(r'^\s*\d+\s*[．.、]?\s*参考文献')


# ---------- 底层工具 ----------
def set_run_font(run, east='宋体', latin=LATIN_FONT, size_pt=None, bold=None, underline=None):
    """设置 run 的中英文字体、字号、加粗、下划线。"""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), latin)
    rFonts.set(qn('w:hAnsi'), latin)
    rFonts.set(qn('w:eastAsia'), east)
    rFonts.set(qn('w:cs'), latin)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    if underline is not None:
        run.font.underline = underline


def set_para_spacing(p, line=LINE_TWIPS, rule='exact', before_lines=None, after_lines=None,
                     before_pt=None, after_pt=None):
    """设置段落的行距与段前/段后。before_lines/after_lines 单位是"行"(0.5=半行)。"""
    pPr = p._element.get_or_add_pPr()
    sp = pPr.find(qn('w:spacing'))
    if sp is None:
        sp = OxmlElement('w:spacing')
        pPr.append(sp)
    if line is not None:
        sp.set(qn('w:line'), str(int(line)))
        sp.set(qn('w:lineRule'), rule)
    # 清除旧的 before/after/行单位，再按需设置
    for attr in ('w:before', 'w:after', 'w:beforeLines', 'w:afterLines'):
        sp.attrib.pop(qn(attr), None)
    if before_lines is not None:
        sp.set(qn('w:beforeLines'), str(int(round(before_lines * 100))))
        sp.set(qn('w:before'), str(int(round(before_lines * 240))))
    if after_lines is not None:
        sp.set(qn('w:afterLines'), str(int(round(after_lines * 100))))
        sp.set(qn('w:after'), str(int(round(after_lines * 240))))
    if before_pt is not None:
        sp.set(qn('w:before'), str(int(before_pt * 20)))
    if after_pt is not None:
        sp.set(qn('w:after'), str(int(after_pt * 20)))


def set_para_indent(p, first_line_twips=None, first_line_chars=None,
                    left_twips=None, left_chars=None, hanging_twips=None):
    """设置首行缩进 / 左缩进 / 悬挂缩进。传 None 表示清空对应项。"""
    pPr = p._element.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    for attr in ('w:firstLine', 'w:firstLineChars', 'w:left', 'w:leftChars', 'w:hanging'):
        ind.attrib.pop(qn(attr), None)
    if first_line_twips is not None:
        ind.set(qn('w:firstLine'), str(int(first_line_twips)))
    if first_line_chars is not None:
        ind.set(qn('w:firstLineChars'), str(int(first_line_chars)))
    if left_twips is not None:
        ind.set(qn('w:left'), str(int(left_twips)))
    if left_chars is not None:
        ind.set(qn('w:leftChars'), str(int(left_chars)))
    if hanging_twips is not None:
        ind.set(qn('w:hanging'), str(int(hanging_twips)))


def set_snap_to_grid(p, val='0'):
    pPr = p._element.get_or_add_pPr()
    el = pPr.find(qn('w:snapToGrid'))
    if el is None:
        el = OxmlElement('w:snapToGrid')
        pPr.append(el)
    el.set(qn('w:val'), val)


def set_tab_stop(p, pos_twips, align=WD_TAB_ALIGNMENT.LEFT):
    """在段落设置一个制表位（用于封面填充区对齐）。"""
    pPr = p._element.get_or_add_pPr()
    tabs = pPr.find(qn('w:tabs'))
    if tabs is None:
        tabs = OxmlElement('w:tabs')
        pPr.append(tabs)
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'left')
    tab.set(qn('w:pos'), str(int(pos_twips)))
    tabs.append(tab)


# ---------- 分类 ----------
def classify(text, style_name):
    """返回 1=节 2=条 3=款 0=正文 None=封面行。优先用 Word 内置 Heading 样式，再回退正则。"""
    if style_name:
        m = re.search(r'Heading\s*([1-9])', style_name)
        if m:
            return min(int(m.group(1)), 3)
    t = text.strip()
    if not t:
        return None
    if RE_L3.match(t):
        return 3
    if RE_L2.match(t):
        return 2
    if RE_L1.match(t):
        return 1
    # （1）（2）… 等项：按正文处理（首行缩进2字符），不再单独作为标题
    return 0


def is_cover_line(text):
    for _, pat in COVER_PATTERNS:
        if pat.search(text):
            return True
    return False


def split_cover_fill(text):
    """把封面填表行切分为 (标签文本, 填充值文本)。"""
    for pat, label in FILL_LABELS:
        m = pat.search(text)
        if m:
            idx = m.end()
            return text[:idx], text[idx:]
    return text, ''


# ---------- 应用格式 ----------
def apply_level(p, level, in_ref=False):
    """对单个段落套用格式。level: 1..3 标题, 0 正文, 'cover' 封面。"""
    text = p.text
    pf = p.paragraph_format
    pf.alignment = None
    set_snap_to_grid(p, '0')

    if level == 'cover':
        apply_cover_line(p, text)
        return

    if level in (1, 2, 3):
        spec = LEVELS[level]
        set_para_spacing(p, line=LINE_TWIPS, rule='exact',
                         before_lines=spec['before_lines'], after_lines=spec['after_lines'])
        set_para_indent(p)  # 标题顶格
        for run in p.runs:
            set_run_font(run, east=spec['east'], latin=LATIN_FONT, size_pt=spec['size'],
                         bold=False, underline=False)
    else:  # 正文（含参考文献条目）
        set_para_spacing(p, line=LINE_TWIPS, rule='exact',
                         before_lines=0, after_lines=0)
        if in_ref:
            # 参考文献条目：悬挂缩进（首行顶格，续行缩进2字符）
            set_para_indent(p, left_twips=REF_HANG_TWIPS, hanging_twips=REF_HANG_TWIPS)
        else:
            # 正文：首行缩进 2 字符（含（1）（2）… 项）
            set_para_indent(p, first_line_twips=BODY_FIRST_TWIPS, first_line_chars=200)
        for run in p.runs:
            set_run_font(run, east=BODY['east'], latin=LATIN_FONT, size_pt=BODY['size'],
                         bold=False, underline=False)


def apply_cover_line(p, text):
    """按模板封面格式设置单行。居中行清除缩进；填表行只给填充区加下划线。"""
    set_snap_to_grid(p, '0')
    if re.search(r'哈尔滨工业大学', text):
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_indent(p)
        set_para_spacing(p, line=None, rule='auto')
        for run in p.runs:
            set_run_font(run, east='楷体_GB2312', latin=LATIN_FONT, size_pt=18, bold=True)
    elif re.search(r'硕士学位开题报告', text):
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_indent(p)
        set_para_spacing(p, line=None, rule='auto')
        for run in p.runs:
            set_run_font(run, east='宋体', latin='Times New Roman', size_pt=24, bold=True)
    elif re.search(r'^[（(](学位论文|实践成果)[)）]', text.strip()):
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_indent(p)
        set_para_spacing(p, before_pt=12, after_pt=0)
        for run in p.runs:
            set_run_font(run, east='宋体', latin=LATIN_FONT, size_pt=16, bold=True)
    elif re.search(r'^\s*题\s*目\s*[:：]', text):
        set_para_spacing(p, before_pt=12, after_pt=0)
        for run in p.runs:
            set_run_font(run, east='宋体', latin=LATIN_FONT, size_pt=18, bold=True)
    elif re.search(r'研究生院制', text):
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_indent(p)
        set_para_spacing(p, line=None, rule='auto')
        for run in p.runs:
            set_run_font(run, east='宋体', latin=LATIN_FONT, size_pt=16, bold=True)
    else:  # 学院/学科/导师/研究生/学号/日期 等填表行
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_spacing(p, line=420, rule='auto', before_lines=0, after_lines=0)
        set_para_indent(p, left_twips=COVER_LEFT_TWIPS)
        label, value = split_cover_fill(text)
        # 清空原 run，重写为 标签(不加下划线) + 制表位 + 填充值(加下划线)
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        if label:
            rl = p.add_run(label)
            set_run_font(rl, east='宋体', latin=LATIN_FONT, size_pt=16, bold=True, underline=False)
        set_tab_stop(p, COVER_FILL_TAB_TWIPS)
        rt = p.add_run('\t')
        set_run_font(rt, east='宋体', latin=LATIN_FONT, size_pt=16, bold=True, underline=False)
        val = value if value.strip() else '　' * 8
        rv = p.add_run(val)
        set_run_font(rv, east='宋体', latin=LATIN_FONT, size_pt=16, bold=True, underline=True)


# ---------- 主流程 ----------
def set_page_setup(doc):
    sec = doc.sections[0]
    sec.page_width, sec.page_height = PAGE['width'], PAGE['height']
    sec.top_margin, sec.bottom_margin = PAGE['top'], PAGE['bottom']
    sec.left_margin, sec.right_margin = PAGE['left'], PAGE['right']
    sec.header_distance, sec.footer_distance = PAGE['header'], PAGE['footer']
    # 清除页眉（报告不设页眉）
    for h in (sec.header, sec.first_page_header):
        for p in h.paragraphs:
            for r in list(p.runs):
                r.text = ''


def build_cover(doc, meta):
    """在文档最前面插入模板封面（含若干空行），正文从分页符后开始。"""
    first = doc.paragraphs[0]

    def add_blank():
        first.insert_paragraph_before('')

    def add_centered(text, size, east='宋体'):
        p = first.insert_paragraph_before('')
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_snap_to_grid(p, '0')
        set_para_indent(p)
        r = p.add_run(text)
        set_run_font(r, east=east, latin=LATIN_FONT, size_pt=size, bold=True)

    def add_topic(text):
        p = first.insert_paragraph_before('')
        set_snap_to_grid(p, '0')
        set_para_spacing(p, before_pt=12, after_pt=0)
        r = p.add_run(text)
        set_run_font(r, east='宋体', latin=LATIN_FONT, size_pt=18, bold=True)

    def add_fill(label, value):
        p = first.insert_paragraph_before('')
        set_snap_to_grid(p, '0')
        set_para_spacing(p, line=420, rule='auto', before_lines=0, after_lines=0)
        set_para_indent(p, left_twips=COVER_LEFT_TWIPS)
        rl = p.add_run(label)
        set_run_font(rl, east='宋体', latin=LATIN_FONT, size_pt=16, bold=True, underline=False)
        set_tab_stop(p, COVER_FILL_TAB_TWIPS)
        rt = p.add_run('\t')
        set_run_font(rt, east='宋体', latin=LATIN_FONT, size_pt=16, bold=True, underline=False)
        val = value if value and value.strip() else '　' * 8
        rv = p.add_run(val)
        set_run_font(rv, east='宋体', latin=LATIN_FONT, size_pt=16, bold=True, underline=True)

    add_blank(); add_blank()
    add_centered('哈尔滨工业大学', 18, east='楷体_GB2312')
    add_blank()
    add_centered('硕士学位开题报告', 24)
    add_blank()
    add_topic('题  目：' + meta.get('title', ''))
    degree = meta.get('degree', '学位论文')
    add_centered('（' + ('学位论文' if degree == '学位论文' else '实践成果') + '）', 16)
    add_fill('学     院（部）', meta.get('college', ''))
    add_fill('学科/专业学位类别', meta.get('discipline', ''))
    add_fill('导        师', meta.get('advisor', ''))
    add_fill('研   究   生', meta.get('student', ''))
    add_fill('学       号', meta.get('student_id', ''))
    add_fill('开题报告日期', meta.get('date', ''))
    add_blank(); add_blank(); add_blank()
    add_centered('研究生院制', 16)
    add_blank()
    # 分页：正文另起一页
    pb = first.insert_paragraph_before('')
    pb.add_run().add_break(WD_BREAK.PAGE)


def main():
    ap = argparse.ArgumentParser(description='HIT 硕士开题报告格式调整')
    ap.add_argument('input', help='输入 docx 路径')
    ap.add_argument('output', help='输出 docx 路径')
    ap.add_argument('--cover-json', help='可选：封面元数据 JSON 文件路径')
    ap.add_argument('--dry-run', action='store_true', help='只统计不改写')
    args = ap.parse_args()

    if not os.path.exists(args.input):
        sys.exit('输入文件不存在: %s' % args.input)

    doc = Document(args.input)
    set_page_setup(doc)

    meta = {}
    if args.cover_json:
        with open(args.cover_json, encoding='utf-8') as f:
            meta = json.load(f)

    # 逐段分类并应用格式
    in_ref = False
    stats = {'cover': 0, 'L1': 0, 'L2': 0, 'L3': 0, 'body': 0}
    for i, p in enumerate(doc.paragraphs):
        text = p.text
        stripped = text.strip()
        if not stripped:
            # 空行保持空，不套格式
            continue
        if i < COVER_MAX_INDEX and is_cover_line(text) and not in_ref:
            apply_level(p, 'cover', in_ref=False)
            stats['cover'] += 1
            continue
        level = classify(text, p.style.name)
        # 进入参考文献区
        if REF_HEAD_RE.match(stripped):
            in_ref = True
        elif level in (1,) and in_ref:
            # 下一个节标题，退出参考文献区
            in_ref = False
        apply_level(p, level, in_ref=in_ref)
        key = {0: 'body', 1: 'L1', 2: 'L2', 3: 'L3'}.get(level, 'body')
        stats[key] += 1

    if args.cover_json and not any(is_cover_line(p.text) for p in doc.paragraphs[:10]):
        build_cover(doc, meta)

    if args.dry_run:
        print('dry-run stats:', stats)
        return

    doc.save(args.output)
    print('已输出: %s' % args.output)
    print('统计: %s' % stats)


if __name__ == '__main__':
    main()
